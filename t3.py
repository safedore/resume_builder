from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

name = 'Safedore'
email = 'safedore.private@gmail.com'
phone = '7034507102'
experience = '2'
education = 'Bachelor of Computer Applications'

summary = ("Enthusiastic and confident web developer with 2 years of hands-on experience in Flutter, Flask, "
           "and Django. Proficient in developing high-performance applications and websites. Skilled in optimizing "
           "performance under tight deadlines and committed to delivering exceptional results. Thrives in "
           "collaborative environments and excels under strong leadership.")

skills = "Python, Dart, Django, Flask, Flutter, HTML, JavaScript, Android"

def get_user_info():
    # print("Please enter your information:")
    # name = input("Name: ")
    # email = input("Email: ")
    # phone = input("Phone: ")
    # experience = input("Experience: ")
    # education = input("Education: ")
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "projects": [
            "College Website (Django, Html, Mysql, Javascript, CSS) - A complete college website",
            "Advice Safari (Django, Html, Mysql, Javascript, CSS, Flutter) - A tourism app",
            "Petofia (Flask, Html, Mysql, Javascript, CSS, Android-JAVA) - An online pet accessory shop",
            "Form Assistant (Django, Mysql, Flutter) - An automatic form filling app"
        ]
    }

def add_heading(doc, text, level, align='left'):
    heading = doc.add_heading(text, level=level)
    if align == 'center':
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_paragraph(doc, text, bullet=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    if bullet:
        paragraph.style = 'List Bullet'
    return paragraph

def generate_resume(info):
    doc = Document()
    add_heading(doc, 'Resume', 0, 'center')

    # Personal Information
    add_heading(doc, 'Personal Information', level=1)
    add_paragraph(doc, f"Name: {info['name']}")
    add_paragraph(doc, f"Email: {info['email']}")
    add_paragraph(doc, f"Phone: {info['phone']}")

    # Summary
    add_heading(doc, 'Summary', level=1)
    add_paragraph(doc, info['summary'])

    # Experience
    add_heading(doc, 'Experience', level=1)
    add_paragraph(doc, info['experience'])

    # Education
    add_heading(doc, 'Education', level=1)
    add_paragraph(doc, info['education'])

    # Technical Skills
    add_heading(doc, 'Technical Skills', level=1)
    add_paragraph(doc, info['skills'])

    # Projects
    add_heading(doc, 'Projects', level=1)
    for project in info['projects']:
        add_paragraph(doc, project, bullet=True)

    doc.save('web_developer_resume.docx')
    convert('web_developer_resume.docx')
    print("Resume generated successfully!")

def main():
    user_info = get_user_info()
    generate_resume(user_info)

if __name__ == "__main__":
    main()