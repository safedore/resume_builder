from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

name = 'Safedore'
email = 'safedore.private@gmail.com'
phone = '7034507102'
education = 'Bachelor of Computer Applications'
languages = 'English, Malayalam, Hindi'
linkedin = 'None'
github = 'None'

summary = ("Enthusiastic and confident web developer with 2 years of hands-on experience in Flutter, Flask, "
           "and Django. Proficient in developing high-performance applications and websites. Skilled in optimizing "
           "performance under tight deadlines and committed to delivering exceptional results. Thrives in "
           "collaborative environments and excels under strong leadership.")

skills = "Python", "Dart", "Django", "Flask", "Flutter", "HTML", "JavaScript", "Android"

experience = [
        {
            "company": "Riss Technologies",
            "title": "Web-App Developer",
            "duration": "June 2022 - May 2024",
            "responsibilities": [
                "Developed web applications and instructed students on web-app development using Django/Flask.",
                "Created prototype mobile applications using Flutter.",
                "Collaborated with cross-functional teams to define, design, and ship new features.",
                "Resolved technical issues under tight deadlines.",
                "Acted as a mentor at times to colleagues.",
                "Participated in code reviews and provided constructive feedback to peers."
            ]
        }
    ]

def get_user_info():
    print("Please Wait:")
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "languages": languages,
        "projects": [
            "College Website (Django, Html, Mysql, Javascript, CSS) - A complete college website",
            "Advice Safari (Django, Html, Mysql, Javascript, CSS, Flutter) - A tourism app",
            "Petofia (Flask, Html, Mysql, Javascript, CSS, Android-JAVA) - An online pet accessory shop",
            "Form Assistant (Django, Mysql, Flutter) - An automatic form filling app"
        ]
    }

def add_heading(doc, text, level, align='left'):
    heading = doc.add_heading(text, level=level)
    if align == 'center':
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_paragraph(doc, text, bullet=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if bullet:
        paragraph.style = 'List Bullet'
    return paragraph

def generate_resume(info):
    doc = Document()
    # Setting margins to fit content on a single page
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(30)
        section.bottom_margin = Pt(30)
        section.left_margin = Pt(30)
        section.right_margin = Pt(30)

    add_heading(doc, 'Safedore', 0, 'center')

    # Personal Information
    add_heading(doc, 'Personal Information', level=1)
    personal_info = [
        f"Name: {info['name']} || Email: {info['email']}",
        f"Phone: {info['phone']} || LinkedIn: {info['linkedin']}",
        f"GitHub: {info['github']}"
    ]
    for item in personal_info:
        add_paragraph(doc, item)

    # Summary
    add_heading(doc, 'Summary', level=1)
    add_paragraph(doc, info['summary'])

    # Experience
    add_heading(doc, 'Experience', level=1)
    for exp in info['experience']:
        add_paragraph(doc, f"Company: {exp['company']}, \nTitle: {exp['title']}, \nDuration: {exp['duration']}")
        for responsibility in exp['responsibilities']:
            add_paragraph(doc, responsibility, bullet=True)

    # Education
    add_heading(doc, 'Education', level=1)
    add_paragraph(doc, info['education'])

    # Languages
    add_heading(doc, 'Languages', level=1)
    add_paragraph(doc, info['languages'], bullet=True)

    # Technical Skills
    add_heading(doc, 'Technical Skills', level=1)
    for skill in skills:
        add_paragraph(doc, skill, bullet=True)
    # add_paragraph(doc, info['skills'], bullet=True)

    # Projects
    add_heading(doc, 'Projects', level=1)
    for project in info['projects']:
        add_paragraph(doc, project, bullet=True)

    doc.save('web_developer_resume.docx')
    convert('web_developer_resume.docx')
    print("Resume generated successfully!")

def main():
    user_info = get_user_info()
    generate_resume(user_info)

if __name__ == "__main__":
    main()
