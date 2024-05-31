from docx import Document

def get_user_info():
    print("Please enter your information:")
    name = input("Name: ")
    email = input("Email: ")
    phone = input("Phone: ")
    summary = input("Summary: ")
    experience = input("Experience: ")
    education = input("Education: ")
    skills = input("Technical Skills: ")
    projects = input("Projects: ")
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "projects": projects
    }

def generate_resume(info):
    doc = Document()
    doc.add_heading('Resume', 0)

    # Personal Information
    doc.add_heading('Personal Information', level=1)
    doc.add_paragraph(f"Name: {info['name']}")
    doc.add_paragraph(f"Email: {info['email']}")
    doc.add_paragraph(f"Phone: {info['phone']}")

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(info['summary'])

    # Experience
    doc.add_heading('Experience', level=1)
    doc.add_paragraph(info['experience'])

    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph(info['education'])

    # Technical Skills
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph(info['skills'])

    # Projects
    doc.add_heading('Projects', level=1)
    doc.add_paragraph(info['projects'])

    doc.save('web_developer_resume.docx')
    print("Resume generated successfully!")

def main():
    user_info = get_user_info()
    generate_resume(user_info)

if __name__ == "__main__":
    main()
