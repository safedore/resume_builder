from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

name = 'Riswan Abdussalam'
email = 'safedore.private@gmail.com'
phone = '+91 7034507102'
education = ['Bachelor of Computer Applications', 'June 2019 - April 2022']
languages = ['Malayalam', 'Hindi', '', '', 'English']
linkedin = 'http://www.linkedin.com/in/riswan-abdussalam-1222451b5'
github = 'http://www.github.com/safedore'

summary = ("Enthusiastic and confident web developer with 2 years of hands-on experience in Flutter, Flask, "
           "and Django. Proficient in developing high-performance applications and websites. Skilled in optimizing "
           "performance under tight deadlines and committed to delivering exceptional results. Thrives in "
           "collaborative environments and excels under strong leadership.")

skills = ["Dart, Flutter", "HTML", "JavaScript", "PHP, Laravel", "Python, Django, Flask"]

experience = [
    {
        "company": "Riss Technologies",
        "title": "Web-App Developer",
        "duration": "June 2022 - May 2024",
        "responsibilities": [
            "Developed web applications and instructed students on web-app development using Django/Flask.",
            "Created prototype mobile applications using Flutter.",
            "Collaborated with cross-functional teams to define, design, and ship new features.",
            "Resolved technical issues under tight deadlines.",
            "Acted as a mentor at times to colleagues.",
            "Participated in code reviews and provided constructive feedback to peers."
        ]
    }
]


def get_user_info():
    print("Please Wait:")
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "languages": languages,
        "projects": [
            "College Website (Django, Html, Mysql, Javascript, CSS) - A complete college website",
            "Advice Safari (Django, Html, Mysql, Javascript, CSS, Flutter) - A tourism app",
            "Petofia (Flask, Html, Mysql, Javascript, CSS, Android-JAVA) - An online pet accessory shop",
            "Form Assistant (Django, Mysql, Flutter) - An automatic form filling app"
        ]
    }


def add_heading(doc, text, level, align='left'):
    heading = doc.add_heading(text, level=level)
    if align == 'center':
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_paragraph(doc, text, bullet=False, align=False, color=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if bullet:
        paragraph.style = 'List Bullet'
    return paragraph


def generate_resume(info):
    doc = Document()
    # Setting margins to fit content on a single page
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(10)
        section.left_margin = Pt(30)
        section.right_margin = Pt(30)

    add_heading(doc, 'Riswan Abdussalam', 0, 'center')

    # Personal Information
    add_heading(doc, 'Personal Information', level=1)
    personal_info = [
        f"Name: {info['name']} \n"
        f"Phone: {info['phone']} || Email: {info['email']}",
        f"GitHub: {info['github']}\nLinkedIn: {info['linkedin']}"
    ]
    for item in personal_info:
        if 'github' in item.lower():
            # print(item.)
            add_paragraph(doc, item, color=(0, 0, 255))
        else:
            add_paragraph(doc, item)

    # Summary
    add_heading(doc, 'Summary', level=1)
    add_paragraph(doc, info['summary'], align=True)

    # Experience
    add_heading(doc, 'Experience', level=1)
    for exp in info['experience']:
        add_paragraph(doc, f"Company: {exp['company']}, \nTitle: {exp['title']}, \nDuration: {exp['duration']}", align=False)
        for responsibility in exp['responsibilities']:
            add_paragraph(doc, responsibility, bullet=True, align=False)

        # Projects
        add_heading(doc, 'Projects', level=1)
        for project in info['projects']:
            add_paragraph(doc, project, bullet=True, align=False)
    # Education
    add_heading(doc, 'Education', level=1)
    add_paragraph(doc, f"{education[0]}, \nDuration: {education[1]}", align=False)

    # Create a table for Languages and Technical Skills
    add_heading(doc, 'Technical Skills                                                            Languages Spoken', level=1, align='left')
    table = doc.add_table(rows=max(len(info['skills']), len(info['languages'])) + 1, cols=2)
    table.autofit = True

    # Add headings for Skills and Languages, and make them bold
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Populate the table with skills and languages
    for i in range(-1, max(len(info['skills'])-1, len(info['languages'])-1)):
        if i < len(info['skills']):
            table.cell(i + 1, 0).text = info['skills'][i]
        if i < len(info['languages']):
            table.cell(i + 1, 1).text = info['languages'][i]


    doc.save('web_developer_resume_np.docx')
    convert('web_developer_resume_np.docx')
    print("Resume generated successfully!")


def main():
    user_info = get_user_info()
    generate_resume(user_info)


if __name__ == "__main__":
    main()
