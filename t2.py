from docx import Document

summary = "Enthusiastic and confident web developer with 2 years of hands-on experience in Flutter, Flask, " \
          "and Django. Proficient in developing high-performance applications and websites. Skilled in optimizing " \
          "performance under tight deadlines and committed to delivering exceptional results. Thrives in " \
          "collaborative environments and excels under strong leadership. "

skills = "Python, Dart, Django, Flask, Flutter, HTML, JavaScript, Android"
def get_user_info():
    print("Please enter your information:")
    name = input("Name: ")
    email = input("Email: ")
    phone = input("Phone: ")
    # summary = input("Summary: ")
    experience = input("Experience: ")
    education = input("Education: ")
    # skills = input("Technical Skills: ")
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "projects": [
            "College Website (Django, Html, Mysql, Javascript, CSS) - A complete college website",
            "Advice Safari (Django, Html, Mysql, Javascript, CSS, Flutter) - A tourism app",
            "Petofia (Flask, Html, Mysql, Javascript, CSS, Android-JAVA) - An online pet accessory shop",
            "Form Assistant (Django, Mysql, Flutter) - An automatic form filling app"
        ]
    }

def generate_resume(info):
    doc = Document()
    doc.add_heading('Resume', 0)

    # Personal Information
    doc.add_heading('Personal Information', level=1)
    doc.add_paragraph(f"Name: {info['name']}")
    doc.add_paragraph(f"Email: {info['email']}")
    doc.add_paragraph(f"Phone: {info['phone']}")

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(info['summary'])

    # Experience
    doc.add_heading('Experience', level=1)
    doc.add_paragraph(info['experience'])

    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph(info['education'])

    # Technical Skills
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph(info['skills'])

    # Projects
    doc.add_heading('Projects', level=1)
    for i, project in enumerate(info['projects'], start=1):
        doc.add_paragraph(f"{i}. {project}")

    doc.save('web_developer_resume.docx')

    from docx2pdf import convert
    convert('web_developer_resume.docx')
    print("Resume generated successfully!")

def main():
    user_info = get_user_info()
    generate_resume(user_info)

if __name__ == "__main__":
    main()
